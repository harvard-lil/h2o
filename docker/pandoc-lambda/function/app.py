import boto3
import io
import os
import signal
import subprocess
import tempfile
from itertools import zip_longest
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION_START, WD_HEADER_FOOTER_INDEX, WD_ORIENTATION
from docx.styles.style import WD_STYLE_TYPE
from lxml import etree


def lift_footnote(doc, footnotes_part, ref, texts, id, author=False, docx_sections=False):
    if texts is None:
        parent = ref.getparent()
        parent.remove(ref)
        return
    if ref is None:
        return
    id_att = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id"
    custom_mark_att = (
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}customMarkFollows"
    )
    footnote = OxmlElement("w:footnote")
    footnote.attrib[id_att] = id
    doc_insert = OxmlElement("w:footnoteReference")
    doc_insert.attrib[custom_mark_att] = "1"
    doc_insert.attrib[id_att] = id

    # Content
    for t in texts:
        embedded_refs = t.xpath("w:r[*/w:rStyle[starts-with(@w:val,'FootnoteReference')]]/w:t")
        if docx_sections:
            # because you can't add padding to character styles
            if embedded_refs is not None and len(embedded_refs) > 0:
                embedded_refs[0].text = f"{embedded_refs[0].text}  "
            # remove the leading period/spaces on the first chunk of text— an artifact opf CAP case footnote processing
            # texts lest the footnote marks look really janky
            first_text = t.xpath(
                "w:r/w:t[not(starts-with(., '...')) and (starts-with(., '.') or starts-with(., ' '))]"
            )
            if first_text is not None and len(first_text) > 0:
                first_text[0].text = first_text[0].text.lstrip(". ")

        footnote.append(t)

    # Insert into the footnotes file
    footnotes_part.element.append(footnote)

    # commenting this out but leaving it here for now while the styles are in flux. It adjusts footnote references
    # in-paragraph, in case we need more space, etc.
    # ref.xpath('w:t')[0].text = f"{ref.xpath('w:t')[0].text}"

    # Insert the reference into the doc
    ref.insert(1, doc_insert)


def promote_case_footnotes(doc, docx_sections=False):
    val_att = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
    hyperlink_tag = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink"
    paragraph_tag = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"
    author_footnotes = {}

    for ref in doc.element.xpath("//*[*/w:rStyle[starts-with(@w:val,'FootnoteReference')]]"):
        mark_text = ref.text
        style_node = ref.xpath(".//w:rStyle[starts-with(@w:val,'FootnoteReference')]")[0]
        node_id = style_node.attrib[val_att][18:]
        mark_id = f"{node_id}-{mark_text}"
        if mark_id not in author_footnotes:
            author_footnotes[mark_id] = {
                "id": mark_id,
                "mark": mark_text,
                "refs": [],
                "texts": [],
            }
        style_node.attrib[val_att] = "FootnoteReference"
        author_footnotes[mark_id]["refs"].append(ref)

    for text_el in doc.element.xpath("//w:p[w:pPr/w:pStyle[starts-with(@w:val,'FootnoteText')]]"):
        mark_text = text_el.xpath(".//*[*/w:rStyle]/w:t")[0].text
        node_id = text_el.xpath(".//w:pStyle[starts-with(@w:val, 'FootnoteText')]/@w:val")[0][13:]
        mark_id = f"{node_id}-{mark_text}"
        if mark_id not in author_footnotes:
            author_footnotes[mark_id] = {
                "id": mark_id,
                "mark": mark_text,
                "refs": [],
                "texts": [],
            }
        mark_el = text_el.xpath("./w:pPr/w:pStyle[starts-with(@w:val,'FootnoteText')]")[0]
        mark_el.attrib[val_att] = "FootnoteText"
        for style in text_el.xpath(".//w:pStyle[starts-with(@w:val, 'FootnoteText')]"):
            style.attrib[val_att] = "FootnoteText"
        for style in text_el.xpath(".//w:rStyle[starts-with(@w:val, 'FootnoteRef')]"):
            style.attrib[val_att] = "FootnoteReference"
        author_footnotes[mark_id]["texts"].append([text_el])

    case_footnotes = {}
    for ref in doc.element.xpath("//*[*/w:rStyle[starts-with(@w:val,'CaseFootnoteReference')]]"):
        mark_text = ref.text
        node = ref.xpath(".//w:rStyle[starts-with(@w:val,'CaseFootnoteReference')]")[0]
        node_id = node.attrib[val_att][22:]
        mark_id = f"{node_id}-{mark_text}"
        if mark_id not in case_footnotes:
            case_footnotes[mark_id] = {
                "id": mark_id,
                "mark": mark_text,
                "refs": [],
                "texts": [],
            }
        node.attrib[val_att] = "FootnoteReference"
        parent = ref.getparent()
        gp = parent.getparent()
        # Don't hoist a footnote outside of a paragraph. Word can't handle it.
        if gp is not None and len(gp) and parent.tag == hyperlink_tag and gp.tag == paragraph_tag:
            gp.replace(parent, ref)
        case_footnotes[mark_id]["refs"].append(ref)

    for footnote_start in doc.element.xpath(
        "//*[*/*[starts-with(@w:val,'CaseFootnoteText')] and .//w:hyperlink]"
    ):
        mark_text = footnote_start.xpath(".//*[*/w:rStyle]/w:t")[0].text
        node_id = footnote_start.xpath(
            ".//w:pStyle[starts-with(@w:val, 'CaseFootnoteText')]/@w:val"
        )[0][17:]

        current_stack = [footnote_start]
        next_footnote_candidate = footnote_start.getnext()
        style = next_footnote_candidate.xpath(".//w:pStyle/@w:val")
        link = next_footnote_candidate.xpath(".//w:hyperlink//text()")

        # In cases that a footnote spans multiple paragraphs (judges amiright) roll up those blocks into an array
        while next_footnote_candidate is not None and style and style[0] != "CaseBody" and not link:
            current_stack.append(next_footnote_candidate)
            next_footnote_candidate = next_footnote_candidate.getnext()
            if (
                next_footnote_candidate.tag
                == "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}bookmarkStart"
            ):
                break
            style = next_footnote_candidate.xpath(".//w:pStyle/@w:val")
            if style:
                style_node = next_footnote_candidate.xpath(
                    ".//w:pStyle[starts-with(@w:val,'CaseFootnoteText')]"
                )
                if style_node:
                    style_node[0].attrib[val_att] = "CaseFootnoteText"
            link = next_footnote_candidate.xpath(".//w:hyperlink//text()")
        mark_id = f"{node_id}-{mark_text}"
        if mark_id not in case_footnotes:
            case_footnotes[mark_id] = {
                "id": mark_id,
                "mark": mark_text,
                "refs": [],
                "texts": [],
            }
        hl = footnote_start.getchildren()[1]
        footnote_start.xpath(".//w:rStyle")[0].attrib[val_att] = "FootnoteReference"
        footnote_start.replace(hl, hl.getchildren()[0])
        case_footnotes[mark_id]["texts"].append(current_stack)

    # extract refs and footnotes here.
    fid = 1

    footnote_part = next(f for f in doc.part.package.parts if f.partname == "/word/footnotes.xml")
    footnote_part.element = parse_xml(footnote_part.blob)
    for val in author_footnotes.values():
        for ref, texts in zip_longest(val["refs"], val["texts"]):
            fid += 1
            lift_footnote(
                doc,
                footnote_part,
                ref,
                texts,
                f"{fid}",
                author=True,
                docx_sections=docx_sections,
            )
    for val in case_footnotes.values():
        for ref, texts in zip_longest(val["refs"], val["texts"]):
            fid += 1
            lift_footnote(
                doc,
                footnote_part,
                ref,
                texts,
                f"{fid}",
                author=False,
                docx_sections=docx_sections,
            )
    footnote_part._blob = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        + etree.tostring(footnote_part.element)
    )
    for x in doc.styles.element.xpath("//w:style[starts-with(@w:styleId,'FootnoteText-')]"):
        doc.styles.element.remove(x)
    for x in doc.styles.element.xpath("//w:style[starts-with(@w:styleId,'FootnoteReference-')]"):
        doc.styles.element.remove(x)
    return doc


def sectionizer(
    doc,
    doc_w=8.5,
    doc_h=11,
    internal_margin=1,
    external_margin=1,
    big_margin=1.25,
    gutter=0.5,
):
    """
    Adds section breaks for headers, footers, etc.
    Section breaks control formats affecting more than one paragraph— part of a page, page, chapter, document, etc.
    The myriad formatting options include page number settings,  header and footer definitions, margins, columns,
    page breaks that move to the next odd page rather than just ending the page, and others. In docx XML, they are
    sectPr elements and are either attached to the body itself as the last element— which holds the default
    settings for the entire document— or the last paragraph element of the section.

    We determine where to place sections through paragraph style names. We traverse a list of all paragraph styles
    instead of the xml directly to increase perf. The index in the list matches their index in the doc object— don't
    add new paragraphs mid-doc here without changing this logic.

    SectPrs must be attached to the pPr element in the LAST paragraph in the section (and also one by itself at the
    end of the body element for the document-wide default settings.)

    With columns, each resource/section/chapter header section gets its own, and each of their content blocks get
    their own, too. However, in this simpler configuration, we place:
    * One on the last paragraph of the front matter to apply lower roman page numbering and have blank
      headers/footers, except on the title page where we see placeholder text for the author name.
    * One at the end of each section occuring immediately before a chapter starts. This lets the new chapter have
      always start on an odd page.

    We also normalize the styles of all tables and style lists, here.
    """

    # allows for conversion later. Python-docx uses a different representation of length internally than docx itself,
    # which is in twips, or inches. This will also be useful when we implement scaling based on page size, using
    # non-imperial page sizes, etc.
    doc_w = Inches(doc_w)
    doc_h = Inches(doc_h)
    internal_margin = Inches(internal_margin)
    external_margin = Inches(external_margin)
    big_margin = Inches(big_margin)
    gutter = Inches(gutter)

    # See which rIDs pandoc gave our headers and footers.
    # The rels list holds related xml files like header and footer among other things, and target_ref is the file name.
    rels = doc.part.rels
    headers_and_footers = {
        rels[rid].target_ref.replace(".xml", ""): rels[rid]
        for rid in rels
        if rels[rid].target_ref.endswith(".xml")
        and (rels[rid].target_ref.startswith("header") or rels[rid].target_ref.startswith("footer"))
    }

    body_element = doc.element.xpath("/w:document/w:body")[0]
    doc.settings.element.append(OxmlElement("w:mirrorMargins"))

    def section_break(destination, frontmatter=False, chapter=False, pre_chap=False):
        """Assembles the section break, appends it to the appropriate graf, properly discards the empty wrapper,
        because it gives a hoot and won't pollute."""

        # The doc itself should have an existing default section attached to body, not a paragraph
        if destination == "doc-wide":
            sec = doc.sections[0]
        else:
            sec = doc.add_section()
            destination._element.xpath("w:pPr")[0].append(sec._sectPr)
            body_element.remove(doc.paragraphs[-1]._element)

        if chapter:
            sec.start_type = WD_SECTION_START.ODD_PAGE
        elif pre_chap:
            sec.start_type = WD_SECTION_START.NEW_PAGE
        else:
            sec.start_type = WD_SECTION_START.CONTINUOUS

        sec.bottom_margin = internal_margin
        sec.gutter = gutter
        sec.header_distance = external_margin
        sec.left_margin = external_margin
        sec.orientation = WD_ORIENTATION.PORTRAIT
        sec.page_height = doc_h
        sec.page_width = doc_w
        sec.right_margin = external_margin
        sec.top_margin = big_margin

        sec._sectPr.add_footerReference(
            WD_HEADER_FOOTER_INDEX.EVEN_PAGE, headers_and_footers["footer_blank"].rId
        )
        sec._sectPr.add_footerReference(
            WD_HEADER_FOOTER_INDEX.PRIMARY, headers_and_footers["footer_blank"].rId
        )

        if frontmatter:
            sec.different_first_page_header_footer = True
            sec._sectPr.add_headerReference(
                WD_HEADER_FOOTER_INDEX.EVEN_PAGE,
                headers_and_footers["header_front_matter_even"].rId,
            )
            sec._sectPr.add_headerReference(
                WD_HEADER_FOOTER_INDEX.PRIMARY,
                headers_and_footers["header_front_matter_odd"].rId,
            )
            sec._sectPr.add_footerReference(
                WD_HEADER_FOOTER_INDEX.FIRST_PAGE,
                headers_and_footers["footer_title"].rId,
            )
            sec._sectPr.add_headerReference(
                WD_HEADER_FOOTER_INDEX.FIRST_PAGE,
                headers_and_footers["header_blank"].rId,
            )
            pgNumType_value = "lowerRoman"
        else:
            sec._sectPr.add_headerReference(
                WD_HEADER_FOOTER_INDEX.EVEN_PAGE, headers_and_footers["header_even"].rId
            )
            sec._sectPr.add_headerReference(
                WD_HEADER_FOOTER_INDEX.PRIMARY, headers_and_footers["header_odd"].rId
            )
            pgNumType_value = "decimal"

        # cols value and pgNumType are special cases— they're not supported by python-docx
        pgNumType_query = sec._sectPr.xpath(".//w:pgNumType")
        if len(pgNumType_query) > 0:
            pgNumType_query[0].set(qn("w:fmt"), pgNumType_value)
            if frontmatter:
                pgNumType_query[0].set(qn("w:start"), "0")
            elif len(doc.sections) == 2:
                pgNumType_query[0].set(qn("w:start"), "1")
        else:
            pgNumTypeEl = OxmlElement("w:pgNumType")
            if frontmatter:
                pgNumTypeEl.set(qn("w:start"), "0")
            elif len(doc.sections) == 2:
                pgNumTypeEl.set(qn("w:start"), "1")
            pgNumTypeEl.set(qn("w:pgNumType"), pgNumType_value)
            sec._sectPr.insert_element_before(
                pgNumTypeEl,
                "w:formProt",
                "w:vAlign",
                "w:noEndnote",
                "w:titlePg",
                "w:textDirection",
                "w:bidi",
                "w:rtlGutter",
                "w:docGrid",
                "w:printerSettings",
                "w:sectPrChange",
            )

    our_table_styles = {
        "body": doc.styles.get_by_id("TableText", WD_STYLE_TYPE.PARAGRAPH),
        "headnote": doc.styles.get_by_id("HeadnoteTableText", WD_STYLE_TYPE.PARAGRAPH),
    }

    # much faster than using a clever xpath from the doc root.
    # tables occupying such a small percent of a huge document is a likely culprit.
    for table in doc.tables:
        context_style = (
            doc.tables[0]
            ._element.xpath("preceding-sibling::w:p[1]/w:pPr/w:pStyle")[0]
            .get(qn("w:val"))
        )
        table_style_type = "headnote" if "Headnote" in context_style else "body"
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = our_table_styles[table_style_type]

    # the doc-wide sectpr at the end of the body element
    section_break("doc-wide")

    # Lists in headnotes should have the headnotes style, and lists in body text should have that style.
    # Head End comes before the body, after title, headnotes, etc. Node Start comes before the title, headnote, etc
    our_list_styles = {
        "body": doc.styles.get_by_id("BodyText", WD_STYLE_TYPE.PARAGRAPH),
        "headnote": doc.styles.get_by_id("HeadnoteText", WD_STYLE_TYPE.PARAGRAPH),
    }
    list_style_type = "headnote"
    chapter_head = False
    first_chapter = True
    counter = 0
    for p in doc.paragraphs:
        counter += 1
        style = str(p.style.name)
        if style == "Front Matter End":
            section_break(p, frontmatter=True)
        elif style == "Node Start":
            list_style_type = "headnote"
        elif style == "Head End":
            list_style_type = "body"
            if chapter_head:
                section_break(p, chapter=True)
                chapter_head = False
        elif style == "Chapter Spacer":
            chapter_head = True
            if not first_chapter:
                section_break(p, pre_chap=True)
            first_chapter = False
        elif style == "Compact":
            p.style = our_list_styles[list_style_type]
    return doc


def handler(event, context):

    input_s3_key = event["filename"]
    is_casebook = event["is_casebook"]
    options = event.get("options", {})

    s3_config = {}
    if os.environ.get("USE_S3_CREDENTIALS"):
        s3_config["endpoint_url"] = os.environ["S3_ENDPOINT_URL"]
        s3_config["aws_access_key_id"] = os.environ["AWS_ACCESS_KEY_ID"]
        s3_config["aws_secret_access_key"] = os.environ["AWS_SECRET_ACCESS_KEY"]

    with tempfile.NamedTemporaryFile(suffix=".docx") as pandoc_in:
        # get the source html
        s3 = boto3.resource("s3", **s3_config)
        s3.Bucket(os.environ["EXPORT_BUCKET"]).download_fileobj(input_s3_key, pandoc_in)
        pandoc_in.seek(0)

        # convert to docx with pandoc
        with tempfile.NamedTemporaryFile(suffix=".docx") as pandoc_out:
            command = [
                "pandoc",
                "--from",
                "html",
                "--to",
                "docx",
                "--reference-doc",
                "reference.docx",
                "--output",
                pandoc_out.name,
                "--quiet",
            ]
            if is_casebook:
                command.extend(
                    [
                        "--lua-filter",
                        "table_of_contents.lua"
                    ]
                )
            try:
                response = subprocess.run(
                    command,
                    input=pandoc_in.read(),
                    stderr=subprocess.PIPE,
                    stdout=subprocess.PIPE,
                )
            except subprocess.CalledProcessError as e:
                raise Exception(f"Pandoc command failed: {e.stderr[:100]}")
            if response.stderr:
                raise Exception(f"Pandoc reported error: {response.stderr[:100]}")
            try:
                response.check_returncode()
            except subprocess.CalledProcessError as e:
                if e.returncode < 0:
                    try:
                        sig_string = str(signal.Signals(-e.returncode))
                    except ValueError:
                        sig_string = f"unknown signal {-e.returncode}"
                else:
                    sig_string = f"non-zero exit status {e.returncode}"
                ss = "Pandoc command exited with " + str(sig_string)
                print(ss)
                pandoc_out.seek(0, 0)
                print(response.stderr)
                raise Exception(ss)
            if not os.path.getsize(pandoc_out.name) > 0:
                raise Exception("Pandoc produced no output.")

            doc = Document(pandoc_out)
            if options.get("word_footnotes", False):
                promote_case_footnotes(doc, docx_sections=True)
            sectionizer(doc)
            output = io.BytesIO()
            doc.save(output)
            output.seek(0, 0)
            return output.read()

            return pandoc_out.read()
