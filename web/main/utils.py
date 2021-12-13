import base64
import boto3
from botocore.config import Config
from botocore.exceptions import BotoCoreError, ClientError as BotoClientError
from copy import deepcopy
from datetime import datetime, date
import difflib
from docx import Document
from docx.enum.section import WD_SECTION
import html as python_html
import json
from lxml import etree, html
import mimetypes
import os
from pyquery import PyQuery
import re
import requests
import signal
import subprocess
import tempfile
from urllib.parse import quote, unquote

from django.contrib.auth.tokens import default_token_generator
from django.http import HttpResponse
from django.conf import settings
from django.core.mail import send_mail
from django.template import Context, RequestContext, engines
from django.urls import reverse
from django.utils.encoding import force_bytes
from django.utils.http import urlsafe_base64_encode
from django.utils.text import get_text_list

from .sanitize import sanitize
from .storages import get_s3_storage

import logging
logger = logging.getLogger(__name__)


class APICommunicationError(Exception):
    pass


def show_debug_toolbar(request):
    """
        Whether to show the Django debug toolbar.
    """
    return bool(settings.DEBUG)


def parse_cap_decision_date(decision_date_text):
    """
        Parse a CAP decision date string into a datetime object.

        >>> assert parse_cap_decision_date('2019-10-27') == date(2019, 10, 27)
        >>> assert parse_cap_decision_date('2019-10') == date(2019, 10, 1)
        >>> assert parse_cap_decision_date('2019') == date(2019, 1, 1)
        >>> assert parse_cap_decision_date('2019-02-29') == date(2019, 2, 1)  # non-existent day of month
        >>> assert parse_cap_decision_date('not a date') is None
    """
    try:
        try:
            return datetime.strptime(decision_date_text, '%Y-%m-%d').date()
        except ValueError as e:

            # if court used an invalid day of month (typically Feb. 29), strip day from date
            if e.args[0] == 'day is out of range for month':
                decision_date_text = decision_date_text.rsplit('-', 1)[0]

            try:
                return datetime.strptime(decision_date_text, '%Y-%m').date()
            except ValueError:
                return datetime.strptime(decision_date_text, '%Y').date()
    except Exception:
        # if for some reason we can't parse the date, just store None
        return None


def looks_like_citation(s):
    """
        Return True if string s looks like a case citation (starts and stops with digits).

        >>> all(looks_like_citation(s) for s in [
        ...     "123 Mass. 456",
        ...     "123-mass-456",
        ...     "123 anything else here 456",
        ... ])
        True
        >>> not any(looks_like_citation(s) for s in [
        ...     "123Mass.456",
        ...     "123 Mass.",
        ... ])
        True
    """
    return bool(re.match(r'\d+(\s+|-).*(\s+|-)\d+$', s))

def looks_like_case_law_link(s):
    return bool(re.match(r'^https?://cite\.case\.law(/[/0-9a-zA-Z_-]*)$', s))


def clone_model_instance(instance, **kwargs):
    clone = deepcopy(instance)
    clone.id = clone.pk = clone.created_at = None
    for k, v in kwargs.items():
        setattr(clone, k, v)
    return clone


def fix_after_rails(message):
    """ Use this to document actions that should be taken after the migration to Python is complete. """
    pass


def re_split_offsets(pattern, s):
    """
        Split a string by regular expression, and return the substrings, the offsets for each separator, and the text
        of each separator. This is useful for setting up annotation test templates. Example:

        >>> assert re_split_offsets(r'[A-Z]', "AaaBbbCccDdd") == (
        ...     ["", "aa", "bb", "cc", "dd"],
        ...     [0, 2, 4, 6],
        ...     ["A", "B", "C", "D"],
        ... )
    """
    parts = re.split(f'({pattern})', s)
    strs = [parts[i] for i in range(0, len(parts), 2)]
    split_strs = [parts[i] for i in range(1, len(parts), 2)]
    split_offsets = [sum(len(s) for s in strs[:i+1]) for i in range(len(strs)-1)]
    return strs, split_offsets, split_strs


def normalize_newlines(html_string):
    r"""
        >>> assert normalize_newlines('<p>Hi\r</p>\r\n') == '<p>Hi\n</p>\n'

        We're doing this for a number of reasons.

        1) Consistent line endings make it easier to detect if a string has meaningfully changed.

        2) In our experience, consistent line endings help ensure annotation offsets are handled
        accurately across libraries and languages. Since Django's admin forms POST \n, but the
        WYSIWYG CKEditor uses \r\n, it's easy to end up with a mix of both in the DB... and newlines
        reportedly can be handled differently by different browsers.

        3) Further, the sanitization library "bleach" converts \r to \n, doubling the newlines,
        and forcing annotation offsets to require updating:
        >>> assert sanitize('<p>hello\r</p>\r\n\r\n<p>there</p>') == '<p>hello\n</p>\n\n<p>there</p>'
    """
    return html_string.replace("\r\n", "\n").replace("\r", "\n")


def strip_trailing_block_level_whitespace(html_string):
    r"""
        >>> assert strip_trailing_block_level_whitespace("<p>foo</p>  \r\n \n <p>bar</p>  ") == "<p>foo</p><p>bar</p>"

        We're doing this because the whitespace is being handled by our annotation-placing javascript
        and our css in an undesirable way, resulting in a change to the rendered paragraph numbers
        and visual anomalies, when the whitespace is within an annotated range.

        fix_after_rails("We'd prefer to take a different approach in the JS and the CSS, instead
        of removing the whitespace, but leave that rewrite for another time.")
    """
    tree = parse_html_fragment(html_string)
    for el in tree.iterdescendants():
        if el.tag in block_level_elements and el.tail and re.match(r'\s+$', el.tail):
            el.tail = ''
    return inner_html(tree)


def parse_html_fragment(html_str):
    """
        Parse an html fragment (one or more tags with optional surrounding text) into an lxml tree
        wrapped in a parent <div>.
    """
    # lxml's fragment_fromstring() throws away whitespace that comes at the start of the string if followed by
    # a tag. Avoid this edgecase by stripping leading whitespace and re-appending to our output:
    initial_spaces = re.match(r'\s+', html_str)
    if initial_spaces:
        initial_spaces = initial_spaces.group(0)
        html_str = html_str[len(initial_spaces):]
    else:
        initial_spaces = ''

    #It's here. This next line is the problem.
    el = html.fragment_fromstring(html_str, create_parent=True)

    if initial_spaces:
        el.text = initial_spaces + (el.text or '')

    return el

def format_footnotes_for_export(html_str):
    """
    Footnotes are stored in TextBlocks as
    `<span class="footnote footnote-ref" data-custom-style="Footnote Reference" id="footnote-${data.id}-ref">${data.mark}</span>`
    `<div class="footnote footnote-body" id="footnote-${data.id}"><p><span class="footnote-label" contenteditable="false">${data.mark}</span>${data.footnote}</p></div>`
    """
    if not html_str:
        return html_str
    pq = PyQuery(html_str)
    pq(".footnote-label").append(". ")
    return pq.outer_html()

fix_after_rails("this is redundant of the BLOCK_LEVEL_ELEMENTS javascript array; one could feed the other once we move the asset pipeline over")
block_level_elements = {
    'address', 'article', 'aside', 'blockquote', 'details', 'dialog', 'dd', 'div', 'dl', 'dt', 'fieldset', 'figcaption',
    'figure', 'footer', 'form', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'header', 'hgroup', 'hr', 'li', 'main', 'nav', 'ol',
    'p', 'pre', 'section', 'table', 'ul'
}
# via https://developer.mozilla.org/en-US/docs/Glossary/empty_element
void_elements = {
    'area', 'base', 'br', 'col', 'embed', 'hr', 'img', 'input', 'keygen', 'link', 'meta', 'param', 'source', 'track', 'wbr'
}

def prefix_ids_hrefs(html_str, prefix):
    if not html_str:
        return html_str
    def prefix_id(_, el):
        original_id = el.attrib['id']
        el.attrib['id'] = f"{prefix}-{original_id}"

    def prefix_href(_, el):

        original_target = el.attrib['href'][1:]
        el.attrib['href'] = f"#{prefix}-{original_target}"

    pq = PyQuery(html_str)
    pq("[id]").each(prefix_id)
    pq("[href^='#']").each(prefix_href)
    return pq.outer_html()


def rich_text_export(html_str, request=None, id_prefix=''):
    if not (html_str and request and id_prefix):
        return html_str
    pq = PyQuery(html_str)

    # Footnote labels have an :after css pseudo element with '.' content
    # To prevent editing of the label in a way that will be lost
    pq(".footnote-label").append(". ")

    # IDs that unique within a document may not be unique within multiple documents
    # so we add a prefix
    def prefix_id(_, el):
        original_id = el.attrib['id']
        el.attrib['id'] = f"{id_prefix}-{original_id}"

    def prefix_href(_, el):
        original_target = el.attrib['href'][1:]
        el.attrib['href'] = f"#{id_prefix}-{original_target}"

    pq("[id]").each(prefix_id)
    pq("[href^='#']").each(prefix_href)

    # Images need to have absolute urls
    def absolute_src(_, el):
        original_src = el.attrib['src']
        el.attrib['src'] = request.build_absolute_uri(original_src)

    pq("img[src]").each(absolute_src)

    # Add Doc styling wrappers to images
    def replace_in_parent(style,el):
        original_html = el.parent().html(method="html")
        src = el.outer_html()
        replacement = f"</p><div data-custom-style='{style}'>{el.outer_html()}</div><p>"
        el.parent().html(original_html.replace(src, replacement))

    for el in pq("img.image-center-large").items():
        replace_in_parent("Image Centered Large", el)
    for el in pq("img.image-center-medium").items():
        replace_in_parent("Image Centered Medium", el)
    for el in pq("img.image-left-medium").items():
        replace_in_parent("Image Left Medium", el)
    for el in pq("img.image-right-medium").items():
        replace_in_parent("Image Right Medium", el)

    # Insert a non-breaking space if the paragraph after an image is empty
    # to prevent it from potentially overlapping with a following image
    pq("div[data-custom-style]+p:empty").text(" \xa0 ")

    return pq.outer_html()


def remove_empty_tags(tree, ignore_tags=void_elements):
    """
        Remove empty child elements from an lxml Element, except for any listed in the ignore_tags set. Example:
            >>> tree = etree.XML('<p>asfd<a><b>asdf<c/>asdf</b></a>asdf<d></d></p>')
            >>> remove_empty_tags(tree)
            >>> etree.tostring(tree)
            b'<p>asfd<a><b>asdfasdf</b></a>asdf</p>'
            >>> tree = etree.XML('<p><a><b><c></c></b></a></p>')
            >>> remove_empty_tags(tree, {'a'})
            >>> etree.tostring(tree)
            b'<p><a/></p>'
    """
    for el in tree.iterdescendants():
        while True:
            if el.tag in ignore_tags or el.text or len(el):
                break
            parent = el.getparent()
            if el.tail:
                prev = el.getprevious()
                if prev is None:
                    parent.text = (parent.text or '') + el.tail
                else:
                    prev.tail = (prev.tail or '') + el.tail
            parent.remove(el)
            if parent == tree:
                break
            el = parent


def inner_html(tree):
    """ Return inner HTML of lxml element """
    return (python_html.escape(tree.text) if tree.text else '') + \
        ''.join([html.tostring(child, encoding=str) for child in tree.iterchildren()])


def elements_equal(e1, e2, ignore={}, ignore_trailing_whitespace=False, tidy_style_attrs=False, exc_class=ValueError):
    """
        Recursively compare two lxml Elements.
        Raise an exception (by default ValueError) if not identical.
        Optionally, ignore trailing whitespace after block elements.
        Optionally, munge "style" attributes for easier comparison.
    """
    if e1.tag != e2.tag:
        raise exc_class(f"e1.tag != e2.tag ({e1.tag} != {e2.tag})")
    e1t = (e1.text and e1.text.replace("\n","").strip()) or ''
    e2t = (e2.text and e2.text.replace("\n","").strip()) or ''
    if e1t != e2t:
        diff = '\n'.join(difflib.ndiff([e1.text or ''], [e2.text or '']))
        raise exc_class(f"e1.text != e2.text:\n{diff}")
    e1tail = (e1.tail or '').strip()
    e2tail = (e2.tail or '').strip()
    if e1tail != e2tail:
        exc = exc_class(f"e1.tail != e2.tail ({e1.tail} != {e2.tail})")
        if ignore_trailing_whitespace:
            if (e1.tail or '').strip() or (e2.tail or '').strip():
                raise exc
        else:
            raise exc
    ignore_attrs = ignore.get('attrs', set()) | ignore.get('tag_attrs', {}).get(e1.tag.rsplit('}', 1)[-1], set())
    e1_attrib = {k:v for k,v in e1.attrib.items() if k not in ignore_attrs}
    e2_attrib = {k:v for k,v in e2.attrib.items() if k not in ignore_attrs}
    if tidy_style_attrs and e1_attrib.get('style'):
        # allow easy comparison of sanitized style tags by removing all spaces and final semicolon
        e1_attrib['style'] = e1_attrib['style'].replace(' ', '').rstrip(';')
        e2_attrib['style'] = e2_attrib['style'].replace(' ', '').rstrip(';')
    if e1_attrib != e2_attrib:
        diff = "\n".join(difflib.Differ().compare([f"{i}: {i}" % i for i in sorted(e1_attrib.items())], [f"{i}: {i}" for i in sorted(e2_attrib.items())]))
        raise exc_class(f"e1.attrib != e2.attrib:\n{diff}")
    s1 = [i for i in e1 if i.tag.rsplit('}', 1)[-1] not in ignore.get('tags', ())]
    s2 = [i for i in e2 if i.tag.rsplit('}', 1)[-1] not in ignore.get('tags', ())]
    if len(s1) != len(s2):
        diff = "\n".join(difflib.Differ().compare([s.tag for s in s1], [s.tag for s in s2]))
        raise exc_class(f"e1 children != e2 children:\n{diff}")
    for c1, c2 in zip(s1, s2):
        elements_equal(c1, c2, ignore, ignore_trailing_whitespace, tidy_style_attrs, exc_class)

    # If you've gotten this far without an exception, the elements are equal
    return True


class StringFileResponse(HttpResponse):
    """
        A response that sets Content-Type and Content-Disposition like Django's FileResponse, but takes a string instead
        of a filelike object. This is needed because uwsgi can't handle BytesIO objects --
        see https://github.com/unbit/uwsgi/issues/1126

        Logic based on django.http.response.FileResponse.set_headers.
    """
    def __init__(self, *args, as_attachment=False, filename='', **kwargs):
        super().__init__(*args, **kwargs)

        # set Content-Type
        if self.get('Content-Type', '').startswith('text/html'):
            if filename:
                content_type, encoding = mimetypes.guess_type(filename)
                # Encoding isn't set to prevent browsers from automatically
                # uncompressing files.
                encoding_map = {
                    'bzip2': 'application/x-bzip',
                    'gzip': 'application/gzip',
                    'xz': 'application/x-xz',
                }
                content_type = encoding_map.get(encoding, content_type)
                self['Content-Type'] = content_type or 'application/octet-stream'
            else:
                self['Content-Type'] = 'application/octet-stream'

        # set Content-Disposition
        if filename:
            disposition = 'attachment' if as_attachment else 'inline'
            try:
                filename.encode('ascii')
                file_expr = f'filename="{filename}"'
            except UnicodeEncodeError:
                file_expr = f"filename*=utf-8''{quote(filename)}"
            self['Content-Disposition'] = f'{disposition}; {file_expr}'
        elif as_attachment:
            self['Content-Disposition'] = 'attachment'


def get_ip_address(request):
    """
        Get user's IP address from request object.
        Use Cloudflare CF-Connecting-IP header, falling back to REMOTE_ADDR for dev.
    """
    return request.META.get('HTTP_CF_CONNECTING_IP', request.META.get('REMOTE_ADDR'))


def render_plaintext_template_to_string(template, context, request=None):
    """
        Render a template to string WITHOUT Django's autoescaping, for
        use with non-HTML templates. Do not use with HTML templates!
    """
    # load the django template engine directly, so that we can
    # pass in a Context/RequestContext object with autoescape=False
    # https://docs.djangoproject.com/en/1.11/topics/templates/#django.template.loader.engines
    #
    # (though render and render_to_string take a "context" kwarg of type dict,
    #  that dict cannot be used to configure autoescape, but only to pass keys/values to the template)
    engine = engines['django'].engine
    if request:
        ctx = RequestContext(request, context, autoescape=False)
    else:
        ctx = Context(context, autoescape=False)
    return engine.get_template(template).render(ctx)


def send_template_email(subject, template, context, from_address, to_addresses):
    context.update({s: getattr(settings, s) for s in settings.TEMPLATE_VISIBLE_SETTINGS})
    email_text = render_plaintext_template_to_string(template, context)
    success_count = send_mail(
        subject,
        email_text,
        from_address,
        to_addresses,
        fail_silently=False
    )
    return success_count


def send_verification_email(request, user):
    # Send verify-email-address email.
    # This uses the forgot-password flow; logic is borrowed from auth_forms.PasswordResetForm.save()
    verify_link = request.build_absolute_uri(reverse('password_reset_confirm', args=[
        urlsafe_base64_encode(force_bytes(user.pk)),
        default_token_generator.make_token(user),
    ]))
    message = "To activate your account, please click the link below or copy it to your web browser.  " \
        f"You will need to create a new password.\n\n{verify_link}"
    send_mail(
        "An H2O account has been created for you",
        message,
        settings.DEFAULT_FROM_EMAIL,
        [user.email_address],
    )


def send_invitation_email(request, receiving_user, casebook):
    # A new user has been invited to collaborate on a casebook
    # Send an email that looks like a verification email, but explains the invitation
    # They need to set their initial password, so this borrows from send_verification_email above
    verify_link = request.build_absolute_uri(reverse('password_reset_confirm', args=[
        urlsafe_base64_encode(force_bytes(receiving_user.pk)),
        default_token_generator.make_token(receiving_user),
    ]))
    inviting_user = request.user
    message = f"""You have been invited by {inviting_user.email_address} to collaborate on a casebook titled {casebook.title}.

You can set up your account by choosing a password at {verify_link}.

Access this casebook directly at {request.build_absolute_uri(casebook.get_absolute_url())} or visit your dashboard at {request.build_absolute_uri("/")} to see all of your casebooks.

To learn more, please visit our help guide at https://about.opencasebook.org/. If you have questions, you can reach us at info@opencasebook.org.


The H2O open casebook team
Harvard Law School Library"""
    send_mail(
        f"{inviting_user.attribution} has invited you to collaborate on a casebook",
        message,
        settings.DEFAULT_FROM_EMAIL,
        [receiving_user.email_address],
    )


def send_collaboration_email(request, receiving_user, casebook):
    # For already existing users, send a notice that they've been invited to collaborate
    inviting_user = request.user
    message = f"""You have been invited by {inviting_user.email_address} to collaborate on a casebook titled {casebook.title}.

Access this casebook directly at { request.build_absolute_uri(casebook.get_absolute_url())} or visit your dashboard at {request.build_absolute_uri("/")} to see all of your casebooks.

To learn more, please visit our help guide at https://about.opencasebook.org/. If you have questions, you can reach us at info@opencasebook.org.

The H2O open casebook team
Harvard Law School Library"""
    send_mail(
        f"{inviting_user.attribution} has invited you to collaborate on a casebook.",
        message,
        settings.DEFAULT_FROM_EMAIL,
        [receiving_user.email_address],
    )


def get_link_title(url):
    file_name_re = re.compile('/([^/]*)(?:[.].{1,4})$')
    last_slug_re = re.compile('/([^/]*)/$')
    file_name = file_name_re.search(url)
    default_title = url
    last_slug = last_slug_re.search(url)
    if file_name and file_name.groups()[0]:
        default_title = unquote(file_name.groups()[0])
    elif last_slug and last_slug.groups()[0]:
        default_title = unquote(last_slug.groups()[0])
    resp = None
    try:
        resp = requests.get(url, verify=False)
    except Exception:
        return default_title
    if not resp or not resp.ok:
        return default_title
    body = PyQuery(resp.content)
    if not body:
        return default_title
    title = body.find('title')
    if not title:
        return default_title
    return title[0].text


def export_via_pandoc(obj, html, file_type):
    export_type = obj.__class__.__name__

    logger.info(f"Exporting {export_type} {obj.id}: launching pandoc subprocess")
    with tempfile.NamedTemporaryFile(suffix=f'.{file_type}') as pandoc_out:
        command = []
        if file_type == 'json':
            command = [
                'pandoc',
                '--from', 'html',
                '--to', 'json',
                '--output', pandoc_out.name,
                '--quiet'
            ]
        else:
            command = [
                'pandoc',
                '--from', 'html',
                '--to', 'docx',
                '--reference-doc', os.path.join(settings.PANDOC_DIR, 'reference.docx'),
                '--output', pandoc_out.name,
                '--quiet',
                '--self-contained'
            ]
        if export_type == 'Casebook':
            command.extend(['--lua-filter', os.path.join(settings.PANDOC_DIR, 'table_of_contents.lua')])

        try:
            response = subprocess.run(command, input=html.encode('utf8'), stderr=subprocess.PIPE,
                                      stdout=subprocess.PIPE)
        except subprocess.CalledProcessError as e:
            if export_type == 'Casebook':
                obj.inc_export_fails()
            raise Exception(f"Pandoc command failed: {e.stderr[:100]}")
        if response.stderr:
            if export_type == 'Casebook':
                obj.inc_export_fails()
            raise Exception(f"Pandoc reported error: {response.stderr[:100]}")
        try:
            response.check_returncode()
        except subprocess.CalledProcessError as e:
            if export_type == 'Casebook':
                obj.inc_export_fails()
            if e.returncode < 0:
                try:
                    sig_string = str(signal.Signals(-e.returncode))
                except ValueError:
                    sig_string = f"unknown signal {-e.returncode}"
            else:
                sig_string = f"non-zero exit status {e.returncode}"
            raise Exception(f"Pandoc command exited with {sig_string}")

        if export_type == 'Casebook' and obj.export_fails > 0:
            obj.reset_export_fails()
        return pandoc_out.read()


def export_via_aws_lambda(obj, html, file_type):
    export_settings = settings.AWS_LAMBDA_EXPORT_SETTINGS
    export_type = obj.__class__.__name__

    logger.info(f"Exporting {export_type} {obj.id}: uploading source")
    storage = get_s3_storage(
        bucket_name=export_settings['bucket_name'],
        config={k:v for k,v in export_settings.items() if k in ['endpoint_url', 'secret_key', 'access_key'] and v}
    )
    with tempfile.NamedTemporaryFile(suffix='.html') as inputfile:
        # temporarily save the html source to s3, where the lambda can access it
        filename = f"{export_type.lower()}-{obj.id}-{inputfile.name.split('/')[-1]}"
        inputfile.write(bytes(html, 'utf-8'))
        inputfile.seek(0)
        storage.save(filename, inputfile)

        # trigger the lambda and wait for the produced file
        try:
            logger.info(f"Exporting {export_type} {obj.id}: triggering lambda")
            if export_settings.get('function_arn'):
                lambda_client = boto3.client(
                    'lambda',
                    export_settings['function_region'],
                    config=Config(read_timeout=settings.AWS_LAMBDA_EXPORT_TIMEOUT),
                    **({'aws_access_key_id': export_settings['access_key'], 'aws_secret_access_key': export_settings['secret_key']} if export_settings['access_key'] else {})
                )
                raw_response = lambda_client.invoke(
                    FunctionName=export_settings['function_name'],
                    LogType='Tail',
                    Payload=bytes(json.dumps({"filename": filename, "is_casebook": export_type == 'Casebook'}), 'utf-8')
                )
                response = {
                    'status_code': raw_response['ResponseMetadata']['HTTPStatusCode'],
                    'headers': raw_response['ResponseMetadata']['HTTPHeaders'],
                    'content': raw_response['Payload'],
                    'get_text': lambda: raw_response['Payload'].read()
                }
                lambda_log_str = str(base64.b64decode(raw_response['LogResult']), 'utf-8').strip().replace('\n', '; ').replace('\t', ', ')
                logger.info(f"Exporting Casebook 4227: Lambda logs \"{lambda_log_str}\"")
            else:
                raw_response = requests.post(
                    export_settings['function_url'],
                    timeout=settings.AWS_LAMBDA_EXPORT_TIMEOUT,
                    json={
                        'filename': filename,
                        'is_casebook': export_type == 'Casebook'
                    }
                )
                response = {
                    'status_code': raw_response.status_code,
                    'headers': {k.lower():v for k,v in raw_response.headers.items()},
                    'log': None,
                    'content': raw_response.content,
                    'get_text': lambda: raw_response.text
                }
            assert response['status_code'] == 200, f"Status: {response['status_code']}. Content: {response['get_text']()}"
            assert not response['headers'].get('x-amz-function-error') and response['headers']['content-type'] in ['application/zip', 'application/octet-stream'], f"x-amz-function-error: {response['headers'].get('x-amz-function-error')}, content-type:{response['headers']['content-type']}, {response['get_text']()}"
        except (BotoCoreError, BotoClientError, requests.RequestException, AssertionError) as e:
            if export_type == 'Casebook':
                obj.inc_export_fails()
            raise Exception(f"AWS Lambda export failed: {str(e)}")
        finally:
            # remove the source html from s3
            storage.delete(filename)

        # return the docx to the user
        if export_type == 'Casebook' and obj.export_fails > 0:
            obj.reset_export_fails()
        return response['content']


def export_via_python_docx(obj, children):

    document = Document(os.path.join(settings.PANDOC_DIR, 'template.docx'))

    def add_section(start_style, vertical_alignment='top'):
        document.add_section(start_style)._sectPr.append(etree.fromstring(f'<w:vAlign w:val="{vertical_alignment}" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))

    def author_string():
        return get_text_list([author.display_name for author in obj.primary_authors], 'and')

    def add_titles(node, node_type):
        if hasattr(node, 'ordinals'):
            document.add_paragraph(node.ordinal_string(), style=f"{node_type} Number")
        document.add_paragraph(node.title, style=f"{node_type} Title")
        if node.subtitle:
            document.add_paragraph(node.subtitle, style=f"{node_type} Subtitle")

    # the first section doesn't need to be added; you are already in the first section.
    # so, if you immediately add an WD_SECTION.ODD_PAGE section, the imaginary cursor is
    # writing on page 3.

    # the first doc section is an H2O preamble, with instructions
    document.add_paragraph('(Placeholder for the preamble)')

    # the next, if we want one, is a half-title page, but I think we don't
    # the next is the title page
    add_section(WD_SECTION.ODD_PAGE, 'bottom')
    add_titles(obj, 'Casebook')
    document.add_paragraph(author_string(), style="Casebook Authors")

    # on the reverse is copyright info
    add_section(WD_SECTION.EVEN_PAGE, 'bottom')
    document.add_paragraph(f"© {author_string()}")
    document.add_paragraph("This work is licensed to the public under a Creative Commons Attribution NonCommercial-Share Alike 3.0 license (international):")
    document.add_paragraph("http://creativecommons.org/licenses/by-nc-sa/3.0/us/")

    # the next section is the table of contents

    # then we get into the book's contents....
    #....the first item of which is probably the casebook's headnote.
    #....in the HTML version, that is displayed above the TOC,
    #....but that's very unusual e.g., https://www.bookcoverdesigner.com/book-interior-content/
    #....So, we probably need to ask authors how they want that text treated during export,
    #....but might default to a Preface, just after the TOC.
    #....With this casebook, it's a sub-subtitle (https://opencasebook.org/casebooks/523-advanced-constitutional-law/)
    #....What else might make sense? let's look at some casebook headnotes and see how they are used.

    # but okay now we get into the actual contents
    for child in children:
        child_type = child.get_export_class()

        # ADD THE SECTION
        if child_type in ['Chapter', 'Leading Resource']:
            start_type = WD_SECTION.ODD_PAGE
        elif child_type == 'Section':
            start_type = WD_SECTION.NEW_PAGE
        elif child_type == 'Subsection':
            start_type = WD_SECTION.CONTINUOUS
        else:
            start_type = WD_SECTION.CONTINUOUS
        add_section(start_type)

        # CONFIGURE ITS PAGE HEADER
        # There are three header properties on Section: .header, .even_page_header, and .first_page_header
        # Any existing even page header definitions are preserved when .odd_and_even_pages_header_footer is False; they are simply not rendered by Word. Assigning True to .odd_and_even_pages_header_footer does not automatically create new even header definition
        # Assigning True to .different_first_page_header_footer does not automatically create a new first page header definition
        # https://python-docx.readthedocs.io/en/latest/dev/analysis/features/header.html?highlight=even#header-and-footer

        # CONFIGURE ITS PAGE NUMBERS
        # This is where we can specify if its pages numbers should be arabic, roman, etc.,
        # and whether the counting should start fresh or should be contiguous with the last section.

        # ADD ITS TITLE
        # (This is also where we will need to place the bookmark for the TOC)
        add_titles(child, child_type)

        # ADD ITS CONTENT

        # First, again, goes author-provided headnotes:
        # - page break or not? probably not.
        # - how do we do the page headers, if this is in a chapter, and is multi-page?

        # Then, if the node has it's own content, it goes here, potentially with annotations and formatted footnotes.
        if child.has_body:
            # for now, just dump some text in, to give the textbook content.
            paragraphs = PyQuery(parse_html_fragment(child.export_content(None))).text().split('\n')
            document.add_paragraph(paragraphs[0], style=f"First Paragraph")
            for p in paragraphs[1:]:
                document.add_paragraph(p, style=f"Body Text")


    # and finally, the book's endmatter, which right now is just credits

    # save and return
    with tempfile.NamedTemporaryFile(suffix='.docx') as tmp:
        document.save(tmp)
        tmp.seek(0)
        return tmp.read()
